from __future__ import annotations

import argparse
import html
import json
import re
import sys
import zipfile
from dataclasses import asdict, dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
TEXT_EXTENSIONS = {".txt", ".md", ".csv"}
HTML_EXTENSIONS = {".html", ".htm"}


@dataclass
class Block:
    block_id: str
    source: str
    text: str
    layout_type: str
    confidence: float | None
    extraction_method: str
    page: int | None = None
    section: str | None = None


class SimpleHTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self.skip_depth += 1
        if tag in {"p", "div", "section", "article", "br", "li", "tr", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self.skip_depth:
            self.skip_depth -= 1
        if tag in {"p", "div", "section", "article", "li", "tr", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        return normalize_text(html.unescape(" ".join(self.parts)))


def normalize_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n[ \t]+", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def split_blocks(text: str) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    if len(paragraphs) == 1 and len(paragraphs[0]) > 1200:
        sentences = re.split(r"(?<=[.!?。！？])\s+", paragraphs[0])
        chunks: list[str] = []
        current = ""
        for sentence in sentences:
            if len(current) + len(sentence) > 900 and current:
                chunks.append(current.strip())
                current = sentence
            else:
                current = f"{current} {sentence}".strip()
        if current:
            chunks.append(current.strip())
        return chunks
    return paragraphs


def make_blocks(
    source: Path,
    texts: Iterable[str],
    method: str,
    layout_type: str,
    page: int | None = None,
    confidence: float | None = None,
    prefix: str = "b",
) -> list[Block]:
    blocks: list[Block] = []
    for index, text in enumerate(texts, start=1):
        for sub_index, part in enumerate(split_blocks(text), start=1):
            if not part:
                continue
            page_part = f"p{page}-" if page is not None else ""
            suffix = f"{index}" if sub_index == 1 else f"{index}-{sub_index}"
            blocks.append(
                Block(
                    block_id=f"{page_part}{prefix}{suffix}",
                    source=str(source),
                    page=page,
                    section=None,
                    text=part,
                    layout_type=layout_type,
                    confidence=confidence,
                    extraction_method=method,
                )
            )
    return blocks


def extract_text_file(path: Path, encoding: str) -> list[Block]:
    text = path.read_text(encoding=encoding, errors="replace")
    return make_blocks(path, [text], "python-text", "text")


def extract_html(path: Path, encoding: str) -> list[Block]:
    raw = path.read_text(encoding=encoding, errors="replace")
    parser = SimpleHTMLTextExtractor()
    parser.feed(raw)
    return make_blocks(path, [parser.text()], "python-htmlparser", "html")


def extract_docx(path: Path) -> list[Block]:
    try:
        from docx import Document  # type: ignore

        doc = Document(path)
        texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append(row_text)
        return make_blocks(path, texts, "python-docx", "docx")
    except ImportError:
        return extract_docx_zip(path)


def extract_docx_zip(path: Path) -> list[Block]:
    # Minimal fallback for classroom machines without python-docx.
    with zipfile.ZipFile(path) as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs: list[str] = []
    for paragraph in root.iter(namespace + "p"):
        text = "".join(node.text or "" for node in paragraph.iter(namespace + "t"))
        if text.strip():
            paragraphs.append(text)
    return make_blocks(path, paragraphs, "docx-zip-xml", "docx")


def extract_pdf(path: Path, ocr_lang: str, dpi: int) -> list[Block]:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PDF extraction requires PyMuPDF. Install with: pip install PyMuPDF") from exc

    doc = fitz.open(path)
    blocks: list[Block] = []
    for page_index, page in enumerate(doc, start=1):
        text = normalize_text(page.get_text("text"))
        if text:
            blocks.extend(make_blocks(path, [text], "pymupdf-text", "pdf_text", page=page_index, confidence=1.0))
            continue

        image = render_pdf_page(page, dpi=dpi)
        ocr_text, confidence = ocr_image(image, lang=ocr_lang)
        blocks.extend(make_blocks(path, [ocr_text], "pymupdf-render+tesseract", "pdf_ocr", page=page_index, confidence=confidence))
    return blocks


def render_pdf_page(page: object, dpi: int):
    zoom = dpi / 72
    matrix = page.parent.__class__.__module__  # keeps static analyzers quiet when fitz is optional
    del matrix
    import fitz  # type: ignore
    from PIL import Image  # type: ignore

    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)


def extract_image(path: Path, ocr_lang: str) -> list[Block]:
    try:
        from PIL import Image  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Image extraction requires Pillow. Install with: pip install Pillow") from exc

    image = Image.open(path)
    text, confidence = ocr_image(image, lang=ocr_lang)
    return make_blocks(path, [text], "tesseract", "image_ocr", confidence=confidence)


def ocr_image(image: object, lang: str) -> tuple[str, float | None]:
    try:
        import pytesseract  # type: ignore
    except ImportError as exc:
        raise RuntimeError("OCR requires pytesseract and a local Tesseract install.") from exc

    text = pytesseract.image_to_string(image, lang=lang)
    confidence = None
    try:
        data = pytesseract.image_to_data(image, lang=lang, output_type=pytesseract.Output.DICT)
        values = [float(value) for value in data.get("conf", []) if str(value).strip() not in {"", "-1"}]
        if values:
            confidence = round(sum(values) / len(values) / 100, 4)
    except Exception:
        confidence = None
    return normalize_text(text), confidence


def extract(path: Path, encoding: str, ocr_lang: str, dpi: int) -> list[Block]:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return extract_text_file(path, encoding)
    if suffix in HTML_EXTENSIONS:
        return extract_html(path, encoding)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path, ocr_lang=ocr_lang, dpi=dpi)
    if suffix in IMAGE_EXTENSIONS:
        return extract_image(path, ocr_lang=ocr_lang)
    raise RuntimeError(f"Unsupported file type: {suffix}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract source documents into ontology annotation blocks.")
    parser.add_argument("source", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--document-id", default=None)
    parser.add_argument("--encoding", default="utf-8")
    parser.add_argument("--ocr-lang", default="eng")
    parser.add_argument("--dpi", type=int, default=300)
    args = parser.parse_args()

    if not args.source.exists():
        print(f"Source not found: {args.source}", file=sys.stderr)
        return 2

    blocks = extract(args.source, encoding=args.encoding, ocr_lang=args.ocr_lang, dpi=args.dpi)
    profile = {
        "document_id": args.document_id or args.source.stem,
        "source": str(args.source),
        "block_count": len(blocks),
        "blocks": [asdict(block) for block in blocks],
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {args.output} with {len(blocks)} blocks")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
